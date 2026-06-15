from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from config import MAX_EXTRACTED_TEXT_CHARS
from extractors.base import BaseExtractor
from models import ContentExtractResult


class DocxExtractor(BaseExtractor):
    supported_extensions = {"docx"}

    def extract(self, path: Path) -> ContentExtractResult:
        document = Document(path)
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
            if sum(len(part) for part in parts) >= MAX_EXTRACTED_TEXT_CHARS:
                return ContentExtractResult("\n".join(parts)[:MAX_EXTRACTED_TEXT_CHARS], "Word")
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
                if sum(len(part) for part in parts) >= MAX_EXTRACTED_TEXT_CHARS:
                    return ContentExtractResult("\n".join(parts)[:MAX_EXTRACTED_TEXT_CHARS], "Word")
        return ContentExtractResult("\n".join(parts)[:MAX_EXTRACTED_TEXT_CHARS], "Word")


class XlsxExtractor(BaseExtractor):
    supported_extensions = {"xlsx"}

    def extract(self, path: Path) -> ContentExtractResult:
        workbook = load_workbook(path, read_only=True, data_only=True)
        lines: list[str] = []
        try:
            for sheet in workbook.worksheets:
                for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                    values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                    if values:
                        lines.append(f"[{sheet.title} 行{row_number}] " + " ".join(values))
                    if sum(len(line) for line in lines) >= MAX_EXTRACTED_TEXT_CHARS:
                        return ContentExtractResult("\n".join(lines)[:MAX_EXTRACTED_TEXT_CHARS], "Excel")
        finally:
            workbook.close()
        return ContentExtractResult("\n".join(lines)[:MAX_EXTRACTED_TEXT_CHARS], "Excel")
